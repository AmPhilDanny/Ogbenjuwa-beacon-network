"""
Ogbenjuwa Community Safety Network
Project Documentation Generator
Generates a comprehensive Word (.docx) file for hackathon presentation
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_para(doc, text, bold=False, italic=False, color=None, size=11, indent=False, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table_row(table, cells, header=False, bg=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        cell.paragraphs[0].runs[0].bold = header
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if bg:
            set_cell_bg(cell, bg)
    return row

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C8102E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)
    return p

# ─────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ═══════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('OGBENJUWA')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(200, 16, 46)  # Deep red

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Community Safety Network')
run2.bold = True
run2.font.size = Pt(22)
run2.font.color.rgb = RGBColor(30, 30, 30)

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = tagline.add_run('Protecting Idoma Communities in Benue State, Nigeria')
run3.italic = True
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

# Info box
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Project Type: ').bold = True
info_run = info.add_run('Community Safety & Emergency Response Platform\n')
info_run.font.size = Pt(12)

info.add_run('Coverage: ').bold = True
info_run2 = info.add_run('9 Local Government Areas · Idoma Region · Benue State · Nigeria\n')
info_run2.font.size = Pt(12)

info.add_run('Version: ').bold = True
info_run3 = info.add_run('1.0 — MVP Release · July 2026\n')
info_run3.font.size = Pt(12)

info.add_run('GitHub: ').bold = True
info_run4 = info.add_run('github.com/AmPhilDanny/Ogbenjuwa-beacon-network')
info_run4.font.size = Pt(12)
info_run4.font.color.rgb = RGBColor(0, 70, 180)

doc.add_paragraph()
doc.add_paragraph()

urls = doc.add_paragraph()
urls.alignment = WD_ALIGN_PARAGRAPH.CENTER
urls.add_run('🌐 ogbenjuwa.com.ng  |  🔧 central-command.ogbenjuwa.com.ng  |  👤 portal.ogbenjuwa.com.ng')
for run in urls.runs:
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(200, 16, 46)

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════
add_heading(doc, '1. Executive Summary', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc,
    'Ogbenjuwa (formerly "Amua") is a full-stack community safety alert platform built specifically '
    'for the Idoma-speaking communities of Benue State, Nigeria. The platform enables citizens, '
    'community leaders, and security operatives to report, track, and respond to security incidents '
    'in real time — via both modern web interfaces and traditional SMS/USSD channels (*347#).',
    size=11)

add_para(doc, '')

add_para(doc,
    'In a region facing persistent threats of armed attacks, kidnappings, and communal violence, '
    'Ogbenjuwa bridges the critical information gap between citizens on the ground and coordinated '
    'emergency response. It works on feature phones with 2G internet, supports the Idoma language, '
    'and is designed for the reality of life in rural Benue State.',
    size=11)

add_para(doc, '')
add_para(doc, 'Key Highlights:', bold=True, size=12)
add_bullet(doc, '3 interconnected web applications serving different user groups')
add_bullet(doc, '9 Idoma LGAs covered across Benue State with 934+ mapped villages')
add_bullet(doc, '5 alert types: Armed Attack, Fire, Medical Emergency, Abduction, Other')
add_bullet(doc, '64 admin features across 13 feature categories')
add_bullet(doc, 'Real-time alerts via WebSocket + SMS/USSD (*347#) for feature phones')
add_bullet(doc, 'Bilingual platform: English and Idoma language support')
add_bullet(doc, 'Offline-first citizen app — works on 2G with minimal data')
add_bullet(doc, 'End-to-end deployment on Vercel + Render + Supabase PostgreSQL')

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 2 — PROBLEM STATEMENT
# ═══════════════════════════════════════════════
add_heading(doc, '2. Problem Statement', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc,
    'Idoma communities in Benue State have faced escalating security challenges including armed '
    'herdsmen attacks, kidnappings, and inter-communal violence. The existing emergency response '
    'infrastructure suffers from three critical failures:',
    size=11)

add_para(doc, '')

problems = [
    ('No Real-Time Alert System', 'Residents rely on word-of-mouth, phone calls, or physical bell ringing to spread emergency information — slow, unreliable, and geographically limited.'),
    ('No Coordination Layer', 'Community security operatives (vigilante groups, LGA coordinators) lack a unified command platform. They operate in silos with no shared situational awareness.'),
    ('Digital Divide', 'Most community members do not have smartphones or reliable internet. Any solution must work on basic phones via SMS/USSD.'),
    ('Language Barrier', 'Most platforms are English-only. Elderly residents and low-literacy populations are excluded from digital safety tools.'),
    ('No Accountability', 'There is no system to track patrol activities, log incidents with evidence, or generate analytics to improve community safety planning.'),
]

for title, detail in problems:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(detail).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 3 — SOLUTION
# ═══════════════════════════════════════════════
add_heading(doc, '3. Solution Overview', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc,
    'Ogbenjuwa is a three-layer community safety ecosystem. Each layer serves a distinct user '
    'group with dedicated functionality, all powered by a single unified API backend:',
    size=11)

doc.add_paragraph()

# Solution Layers Table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr = table.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Platform'
hdr[2].text = 'Users'
for cell in hdr:
    set_cell_bg(cell, 'C8102E')
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(11)

rows_data = [
    ('Layer 1 — Beacon Network', 'Public Alert Map\n(ogbenjuwa.com.ng)', 'Any community member viewing the map, reporting incidents, or checking alerts'),
    ('Layer 2 — Citizen Portal', 'Personal Safety App\n(portal.ogbenjuwa.com.ng)', 'Registered citizens who need panic button, personal dashboard, family reunification'),
    ('Layer 3 — Central Command', 'Admin Dashboard + API\n(central-command.ogbenjuwa.com.ng)', 'LGA coordinators, vigilante leaders, super admins managing all operations'),
]

for i, (l, p, u) in enumerate(rows_data):
    row = table.add_row()
    row.cells[0].text = l
    row.cells[1].text = p
    row.cells[2].text = u
    bg = 'FFF5F5' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 4 — ARCHITECTURE
# ═══════════════════════════════════════════════
add_heading(doc, '4. System Architecture', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc, 'The platform uses a microservices-inspired architecture with three independent frontend applications sharing a single Express.js API gateway:', size=11)
doc.add_paragraph()

add_para(doc, 'Architecture Diagram:', bold=True, size=11)
arch = doc.add_paragraph()
arch.paragraph_format.left_indent = Inches(0.3)
arch_run = arch.add_run(
    '┌─────────────────────┐    ┌──────────────────────┐\n'
    '│  Beacon Network     │    │  Citizen Portal       │\n'
    '│  React + Vite       │    │  React + Vite         │\n'
    '│  Port 3000          │    │  Port 3001            │\n'
    '└──────────┬──────────┘    └──────────┬────────────┘\n'
    '           │                          │\n'
    '           └─────────────┬────────────┘\n'
    '                         │  HTTP / JSON / WebSocket\n'
    '               ┌─────────▼────────┐\n'
    '               │  Central Command  │\n'
    '               │  Express API      │\n'
    '               │  Port 4001        │\n'
    '               │  Admin UI: 4000   │\n'
    '               └─────────┬────────┘\n'
    '                         │\n'
    '               ┌─────────▼────────┐\n'
    '               │  PostgreSQL       │\n'
    '               │  (Neon/Supabase) │\n'
    '               └──────────────────┘'
)
arch_run.font.name = 'Courier New'
arch_run.font.size = Pt(9)

doc.add_paragraph()

add_para(doc, 'Frontend Stack (All Apps):', bold=True, size=11)
fe_stack = [
    'Framework: React 18 + TypeScript (strict mode)',
    'Build Tool: Vite 5',
    'Styling: Tailwind CSS 3 + Shadcn UI components',
    'Routing: React Router v6',
    'Maps: Leaflet + react-leaflet (OpenStreetMap tiles)',
    'HTTP Client: Axios with JWT interceptors',
    'Icons: Lucide React',
    'Animations: Framer Motion',
]
for item in fe_stack:
    add_bullet(doc, item)

doc.add_paragraph()

add_para(doc, 'Backend Stack (Central Command):', bold=True, size=11)
be_stack = [
    'Runtime: Node.js + Express.js + TypeScript',
    'ORM: Drizzle ORM (type-safe SQL)',
    'Database: PostgreSQL via Neon (serverless) / Supabase',
    'Authentication: JWT (access tokens 15min + refresh tokens 7 days)',
    'Real-time: WebSocket server (co-located with Express)',
    'SMS/USSD: Africa\'s Talking SDK',
    'Validation: Zod (runtime schema validation)',
    'Container: Docker (for production deployment on Render)',
]
for item in be_stack:
    add_bullet(doc, item)

doc.add_paragraph()

add_para(doc, 'Deployment Infrastructure:', bold=True, size=11)
deploy_items = [
    'Beacon Network (Frontend): Vercel → ogbenjuwa.com.ng',
    'Citizen Portal (Frontend): Vercel → portal.ogbenjuwa.com.ng',
    'Central Command API: Render (Docker) → api.ogbenjuwa.com.ng',
    'Admin Dashboard: Vercel → central-command.ogbenjuwa.com.ng',
    'Database: Neon/Supabase (serverless PostgreSQL)',
]
for item in deploy_items:
    add_bullet(doc, item)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 5 — LAYER 1: BEACON NETWORK
# ═══════════════════════════════════════════════
add_heading(doc, '5. Layer 1 — Ogbenjuwa Beacon Network (Public Alert Map)', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc, 'URL: ogbenjuwa.com.ng  |  Port: 3000  |  Access: Public (no login required)', bold=True, color=(100, 100, 100), size=10)
doc.add_paragraph()

add_para(doc,
    'The Beacon Network is the primary public interface — an interactive safety map that any '
    'community member can access without registration. It serves as the "emergency broadcast system" '
    'for the Ogbenjuwa network.',
    size=11)

doc.add_paragraph()

add_heading(doc, '5.1 Key Features', 2, color=(150, 10, 30))

features_beacon = [
    ('Interactive Alert Map', 'Leaflet.js-powered map with all 9 Idoma LGA boundaries drawn. Shows real-time alert markers color-coded by incident type. Epicentre visualization with 25km radius when an alert is active. Village flash animation on active alerts.'),
    ('Alert Type System', '5 incident categories with distinct colors: Armed Attack (red #DC2626), Fire Outbreak (orange #EA580C), Medical Emergency (blue #2563EB), Abduction (purple #9333EA), Other (gray #6B7280).'),
    ('Incident Feed', 'Chronological list of all reported incidents with severity badges, location info, and time stamps. Filter by alert type.'),
    ('Report Incident', 'Web form for submitting incidents: type selector, GPS location picker, description, severity level (low/medium/high/critical), and photo upload capability.'),
    ('Family Reunification', 'Registry to search for separated family members during emergencies. Submit missing person reports and check reunion status.'),
    ('Patrol Tracking', 'Visual display of active community patrol teams on the map with dead man\'s switch check-in system.'),
    ('SMS/USSD Guide', 'Instructions for using *347# USSD code and SMS commands for feature phones that cannot access the web app.'),
    ('Command Dashboard', 'Operator-facing analytics dashboard with KPI cards, incident heatmaps, severity bar charts, and response time tracking.'),
    ('User Dashboard & Profile', 'Personalized dashboard for registered users with role-adaptive widgets (KPI row, alert history, quick actions, community stats, recent feed, activity log).'),
    ('Offline Indicator', 'Amber banner when device goes offline, with cached data serving last known state.'),
    ('Landing Page', 'Investor-facing landing page with crisis ticker, platform statistics, problem/solution framing, revenue model, and call-to-action.'),
]

for feat_title, feat_desc in features_beacon:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f'{feat_title}: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(feat_desc).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

add_heading(doc, '5.2 Pages / Routes', 2, color=(150, 10, 30))
pages_table = doc.add_table(rows=1, cols=3)
pages_table.style = 'Table Grid'
hdr = pages_table.rows[0].cells
headers = ['Route', 'Page Name', 'Purpose']
for i, h in enumerate(headers):
    hdr[i].text = h
    set_cell_bg(hdr[i], '2D2D2D')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

pages_data = [
    ('/', 'Landing Page', 'Investor/public-facing introduction with platform overview and statistics'),
    ('/home', 'Home Dashboard', 'Safety status overview for logged-in users'),
    ('/alert', 'Alert Map', 'Interactive Leaflet map with real-time incident markers'),
    ('/feed', 'Incident Feed', 'Chronological list of all reported incidents'),
    ('/report', 'Report Incident', 'Submit a new incident report with location and details'),
    ('/patrol', 'Patrol Map', 'Operator patrol tracking with check-in system'),
    ('/reunify', 'Family Reunification', 'Search and register in the displaced persons registry'),
    ('/neighborhood', 'Neighborhood', 'Ward/cluster structure and community leaders'),
    ('/resources', 'Resources', 'Emergency resources: hospitals, police, shelters'),
    ('/dashboard', 'Command Dashboard', 'Analytics: KPIs, heatmaps, incident charts'),
    ('/user-dashboard', 'User Dashboard', 'Personalized role-adaptive user dashboard'),
    ('/profile', 'User Profile', 'Account settings, notification preferences, security'),
    ('/login', 'Login', 'Phone OTP authentication (any 6-digit code in demo)'),
]
for i, (route, name, purpose) in enumerate(pages_data):
    row = pages_table.add_row()
    row.cells[0].text = route
    row.cells[1].text = name
    row.cells[2].text = purpose
    bg = 'F5F5F5' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 6 — LAYER 2: CITIZEN PORTAL
# ═══════════════════════════════════════════════
add_heading(doc, '6. Layer 2 — Citizen Portal (Personal Safety App)', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc, 'URL: portal.ogbenjuwa.com.ng  |  Port: 3001  |  Access: Self-registration via phone OTP', bold=True, color=(100, 100, 100), size=10)
doc.add_paragraph()

add_para(doc,
    'The Citizen Portal is the personal safety companion for every Idoma community member. '
    'It is designed for people who may never have used a safety app before — working on ₦15,000 '
    'phones with minimal data, usable in darkness, in fear, one-handed, and in the Idoma language.',
    size=11)

doc.add_paragraph()

add_heading(doc, '6.1 Core Citizen Capabilities', 2, color=(150, 10, 30))

citizen_features = [
    ('Panic Button', 'The most critical screen. Entire screen dominated by a large red button (200px min diameter). 3-second countdown with "Hold to confirm" ring animation prevents accidental triggers. On confirm: GPS captured, distress signal sent to emergency contacts + LGA admin. Reference code generated (e.g., AMU-PANIC-3847). "I am safe now" cancel button. SMS fallback: PANIC [name] [village] → 347.'),
    ('Active Alert Feed', 'Full list of active and recent alerts in the citizen\'s subscribed LGAs. Filter tabs by LGA. Active alerts section (red header) + Recent section (last 24 hours). Each card shows: type, village, LGA, time, safety instructions, call button. Supports multiple LGA subscriptions.'),
    ('Report Incident (Offline-first)', 'Step-by-step form: incident type (4 icon buttons), GPS location auto-detect, text description (Idoma or English), optional photo capture, urgency level. Reports queue locally in sessionStorage when offline and sync on reconnect.'),
    ('Find Resources', 'Map and list of nearest shelter, water, medical posts, and food during displacement. Leaflet map with resource pins. Capacity bar with color logic (green/amber/red). "Near me" GPS sorting. USSD fallback: *347*2*[LGA code]#'),
    ('Find Family (Reunification)', '4-tab interface: Search registry, Register Myself, Report Missing Person, My Cases. Idoma-first labels. SMS fallback: FIND ME [name] [age] [village] → 347.'),
    ('Community Feed', 'Verified safety updates from community admins and vigilante leaders. 4 post types: Alert (red), Update (amber), Clear (green), Notice (blue). Citizens can "Note" and share posts but cannot post (operators only). Share generates WhatsApp-ready text.'),
    ('My Profile', 'Personal info editing, emergency contacts management (up to 3), LGA alert subscriptions, notification settings, downloadable safety card with registration number.'),
    ('Bilingual Interface', 'Full Idoma / English language toggle on every page. Default language is Idoma. All UI text, labels, and instructions translated. Idoma phrases include: "Ilu di mma" (All Clear), "Ihe ize ndụ!" (DANGER), "Gbaa Oso" (PANIC/Run!).'),
    ('Offline-First Design', 'Caches last 10 alerts, emergency contacts, recent feed posts, and resource list in sessionStorage. Queued reports sync automatically on reconnect. Connection status banner shown at top when offline.'),
    ('72-Hour Sessions', 'Citizens stay logged in for 72 hours (vs 8-hour admin sessions). No passwords — phone OTP only. OTP-less demo mode: any 6-digit code accepted.'),
]

for feat_title, feat_desc in citizen_features:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f'{feat_title}: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(feat_desc).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(5)

doc.add_paragraph()

add_heading(doc, '6.2 Registration Flow', 2, color=(150, 10, 30))
reg_steps = [
    'Step 1: Enter preferred name + Nigerian phone number (+234)',
    'Step 2: Verify 6-digit OTP (demo: any 6 digits)',
    'Step 3: Select LGA (dropdown — 9 Idoma LGAs)',
    'Step 4: Choose language preference (Idoma default / English)',
    'Step 5: Add up to 3 emergency contacts (optional, can set later)',
    '→ Session created for 72 hours → redirect to home dashboard',
]
for step in reg_steps:
    add_bullet(doc, step)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 7 — LAYER 3: CENTRAL COMMAND
# ═══════════════════════════════════════════════
add_heading(doc, '7. Layer 3 — Central Command Deck (Admin + API Backend)', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc, 'URL: central-command.ogbenjuwa.com.ng  |  API Port: 4001  |  Admin Port: 4000', bold=True, color=(100, 100, 100), size=10)
doc.add_paragraph()

add_para(doc,
    'The Central Command Deck is the unified administrative backend — the "war room" for Ogbenjuwa. '
    'It is the single source of truth for all platform data, the command center for security '
    'operations across all LGAs, and the integration hub connecting the Beacon Network and '
    'Citizen Portal.',
    size=11)

doc.add_paragraph()

add_heading(doc, '7.1 User Roles & Permissions', 2, color=(150, 10, 30))

roles_table = doc.add_table(rows=1, cols=3)
roles_table.style = 'Table Grid'
hdr = roles_table.rows[0].cells
for i, h in enumerate(['Role', 'Scope', 'Description']):
    hdr[i].text = h
    set_cell_bg(hdr[i], 'C8102E')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

roles_data = [
    ('super_admin', 'Global', 'System owner — full access, API key management, audit oversight, system configuration'),
    ('state_observer', 'All LGAs', 'Read-only across entire state — analytics, dashboards, reports, audit logs'),
    ('lga_coordinator', 'Assigned LGA', 'Operational command — manage users, alerts, patrols, incidents within their LGA'),
    ('vigilante_leader', 'Assigned team', 'Tactical — manage patrol team, submit check-ins, create/resolve alerts'),
    ('community_admin', 'Assigned LGA', 'Community-facing — manage announcements, resident relations, broadcast'),
    ('resident', 'Own profile', 'Basic access — directory, reports, personal profile (via Citizen Portal)'),
]
for i, (role, scope, desc) in enumerate(roles_data):
    row = roles_table.add_row()
    row.cells[0].text = role
    row.cells[1].text = scope
    row.cells[2].text = desc
    bg = 'FFF5F5' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

add_heading(doc, '7.2 Feature Categories (64 Total Features)', 2, color=(150, 10, 30))

feature_cats = [
    ('Authentication & Security (5 features)', [
        'F-001: Email & password login with rate limiting (5 attempts/min) and account lockout',
        'F-002: JWT token management — 15-min access + 7-day refresh with rotation',
        'F-003: Role-Based Access Control (RBAC) with route and resource-level guards',
        'F-004: Optional TOTP-based 2FA with recovery codes for high-privilege roles',
        'F-005: Session management — view/revoke active sessions, device tracking',
    ]),
    ('User Management (6 features)', [
        'F-006: User directory — paginated/sortable table with role/LGA/status filters, CSV export',
        'F-007: Create user — form with role assignment, auto-generated temporary password',
        'F-008: Edit user — update profile, role, LGA, avatar, reset password',
        'F-009: Deactivate / Reactivate user with reason and audit trail',
        'F-010: User detail view — full profile, activity timeline, assigned alerts/incidents',
        'F-011: Bulk user operations — CSV import, bulk role assignment, bulk deactivation',
    ]),
    ('LGA & Ward Management (4 features)', [
        'F-012: LGA directory — grid view with per-LGA stats: alerts, patrols, coverage %',
        'F-013: Create/Edit LGA — name, code, state, region, coverage target',
        'F-014: Ward management — create/edit/deactivate wards, bulk creation',
        'F-015: LGA dashboard view — per-region detail with all stats, patrol list, ward list',
    ]),
    ('Alert Command Center (6 features)', [
        'F-016: Alert list — filterable by LGA, severity, status, date range, type',
        'F-017: Create alert — auto-incrementing ID (e.g. AL-20260701-042), assign responder, make public',
        'F-018: Alert detail — full info with timeline, evidence, responder notes, related incidents',
        'F-019: Alert dispatch — notify vigilante leaders via WebSocket + SMS, auto-create incident',
        'F-020: Alert resolution — mark resolved/false alarm/escalated with required notes',
        'F-021: Escalation rules — auto-escalate high severity >4hrs to critical, configurable matrix',
    ]),
    ('Incident Management System (6 features)', [
        'F-022: Incident list — filterable table with quick-assign dropdown',
        'F-023: Create incident — form with evidence upload (photos, videos, documents)',
        'F-024: Incident detail — evidence gallery with lightbox, response timeline, status tracker',
        'F-025: Evidence management — drag-and-drop upload, 10MB limit, preview thumbnails',
        'F-026: Incident assignment — assign/reassign with response time tracking',
        'F-027: Incident resolution — closure notes, resident notification, link to alert',
    ]),
    ('Patrol Operations Hub (6 features)', [
        'F-028: Patrol team management — list teams, create/edit, assign leader and LGA',
        'F-029: Team member management — add/remove members, view patrol history',
        'F-030: Shift scheduling — calendar view, conflict detection, status tracking',
        'F-031: Check-in monitoring — real-time feed, compliance rate, late alerts, export log',
        'F-032: Patrol coverage map — LGA-level heat map (green=good, red=poor coverage)',
        'F-033: Patrol reports — end-of-shift PDF generation',
    ]),
    ('SOS Emergency Response (3 features)', [
        'F-034: SOS signal list — real-time feed, priority-sorted, audible browser alert',
        'F-035: SOS signal detail — trigger info, response timeline, assign responder',
        'F-036: SOS dispatch — auto-dispatch to nearest patrol team, auto-create incident',
    ]),
    ('Communications Hub (3 features)', [
        'F-037: Broadcast announcements — rich text editor, target by LGA/role, schedule future publish',
        'F-038: Internal messaging — send to individual user, inbox with read/unread status',
        'F-039: Notification center — in-app bell, alert types, mark all read, desktop push',
    ]),
    ('Real-time Command Dashboard (6 features)', [
        'F-040: Live stats overview — active patrols, alerts by severity (donut), open incidents, SOS',
        'F-041: Live alert feed — real-time WebSocket scrolling list, click to open detail',
        'F-042: LGA summary cards — per-region stats grid with color-coded status indicators',
        'F-043: Incident trend chart — line chart (30 days), filter by LGA/type, daily/weekly/monthly',
        'F-044: Patrol status monitor — real-time list with green/yellow/red status indicators',
        'F-045: System health monitor — API uptime, DB pool, Redis status, WebSocket count',
    ]),
    ('Analytics & Intelligence (6 features)', [
        'F-046: Incident trends — interactive charts by type, LGA, severity, time of day',
        'F-047: Response time analytics — average by LGA, SLA comparison, breakdown by type',
        'F-048: Patrol performance metrics — coverage %, compliance rate, patrol hours by month',
        'F-049: Resident engagement — registrations, report volume, SOS frequency',
        'F-050: Report export — CSV and PDF, scheduled delivery, executive summary templates',
        'F-051: Data filtering & drill-down — filter presets, full-screen chart mode, click to drill',
    ]),
    ('Integration Gateway (6 features)', [
        'F-052: API key management — generate/revoke keys per layer, usage stats',
        'F-053: Layer 1 integration — sync patrol check-ins, receive alerts from Beacon Network',
        'F-054: Layer 2 integration — sync resident registrations, push public alerts/announcements',
        'F-055: Webhook system — configurable endpoints, retry with exponential backoff, delivery logs',
        'F-056: Rate limiting — per-API-key limits, burst allowance, rate limit headers',
        'F-057: API documentation — OpenAPI 3.0 spec, interactive Swagger UI at /api/docs',
    ]),
    ('Audit & Compliance (3 features)', [
        'F-058: Audit log viewer — full action log: timestamp, user, action, resource, IP, CSV export',
        'F-059: Activity timeline — per-user chronological action history',
        'F-060: Data retention — configurable retention period, auto-archive, export before deletion',
    ]),
    ('System Administration (4 features)', [
        'F-061: System settings — alert escalation rules, coverage targets, password policy, 2FA',
        'F-062: Theme configuration — light/dark mode toggle, brand customization',
        'F-063: Database management — migration status, seed data, backup trigger',
        'F-064: Logs & monitoring — server log viewer, error tracking, performance monitoring',
    ]),
]

for cat_title, cat_features in feature_cats:
    add_para(doc, cat_title, bold=True, color=(200, 16, 46), size=11)
    for f in cat_features:
        add_bullet(doc, f, level=1)
    doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 8 — DATA MODEL
# ═══════════════════════════════════════════════
add_heading(doc, '8. Database & Data Model', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc, 'The platform uses PostgreSQL (Neon/Supabase) with Drizzle ORM. The schema consists of 18 core tables:', size=11)
doc.add_paragraph()

tables_data = [
    ('users', 'Platform users with roles, LGA assignment, contact info, and auth credentials'),
    ('sessions', 'JWT refresh token sessions with device info and expiry tracking'),
    ('lgas', 'Local Government Areas with coverage targets and region metadata'),
    ('wards', 'Sub-LGA ward divisions within each LGA'),
    ('alerts', 'Security alerts with type, severity, location, status, and assignment'),
    ('incidents', 'Detailed incident records linked to alerts with evidence and response notes'),
    ('incident_evidence', 'Photos, videos, and documents uploaded as incident evidence'),
    ('patrol_teams', 'Vigilante patrol teams with leader assignment and LGA scope'),
    ('patrol_members', 'Many-to-many relationship between teams and member users'),
    ('patrol_shifts', 'Scheduled patrol shifts with start/end times and status tracking'),
    ('patrol_checkins', 'GPS check-in records from patrol members during active shifts'),
    ('sos_signals', 'Emergency SOS triggers from citizens with location and response status'),
    ('announcements', 'Public/targeted broadcasts from coordinators and admins'),
    ('messages', 'Internal messages between admin users'),
    ('api_keys', 'API keys for Layer 1/2 integration with scopes and usage tracking'),
    ('audit_logs', 'Complete audit trail of all platform actions with IP and details'),
    ('connections', 'Social network connections between citizen users'),
    ('resident_reports', 'Incident reports submitted by citizens (synced from Layer 2)'),
]

db_table = doc.add_table(rows=1, cols=2)
db_table.style = 'Table Grid'
hdr = db_table.rows[0].cells
for i, h in enumerate(['Table Name', 'Description']):
    hdr[i].text = h
    set_cell_bg(hdr[i], '2D2D2D')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

for i, (name, desc) in enumerate(tables_data):
    row = db_table.add_row()
    row.cells[0].text = name
    row.cells[1].text = desc
    bg = 'F5F5F5' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# Alert Types Table
add_heading(doc, '8.1 Alert Types', 2, color=(150, 10, 30))
alert_table = doc.add_table(rows=1, cols=4)
alert_table.style = 'Table Grid'
hdr = alert_table.rows[0].cells
for i, h in enumerate(['Key', 'Label', 'Icon', 'Color Code']):
    hdr[i].text = h
    set_cell_bg(hdr[i], 'C8102E')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

alert_types = [
    ('attack', 'Armed Attack', '⚔️', '#DC2626 (Red)'),
    ('fire', 'Fire Outbreak', '🔥', '#EA580C (Orange)'),
    ('medical', 'Medical Emergency', '🏥', '#2563EB (Blue)'),
    ('abduction', 'Abduction', '🆘', '#9333EA (Purple)'),
    ('other', 'Other Emergency', '⚠️', '#6B7280 (Gray)'),
]
for i, row_data in enumerate(alert_types):
    row = alert_table.add_row()
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        bg = 'FFF5F5' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[j], bg)
        for para in row.cells[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 9 — API ENDPOINTS
# ═══════════════════════════════════════════════
add_heading(doc, '9. API Endpoints Summary', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc, 'Base URL: https://api.ogbenjuwa.com.ng/api/v1  |  Documentation: /api/docs (Swagger UI)', size=10, color=(100, 100, 100))
doc.add_paragraph()

endpoint_groups = [
    ('Authentication', [
        'POST /auth/login — Email + password → access + refresh tokens',
        'POST /auth/refresh — Refresh access token',
        'POST /auth/logout — Invalidate session',
        'GET  /auth/me — Current user profile',
        'POST /auth/forgot-password / reset-password',
    ]),
    ('Users', [
        'GET  /users — List all users (paginated, filterable)',
        'POST /users — Create user',
        'GET  /users/:id — User detail',
        'PUT  /users/:id — Update user',
        'PATCH /users/:id/deactivate | /reactivate',
    ]),
    ('Alerts', [
        'GET  /alerts — List alerts (filter by LGA, severity, status, type)',
        'POST /alerts — Create alert',
        'GET  /alerts/:id — Alert detail',
        'PATCH /alerts/:id/resolve — Resolve alert',
        'PATCH /alerts/:id/escalate — Escalate severity',
    ]),
    ('Incidents', [
        'GET  /incidents — List incidents',
        'POST /incidents — Create incident',
        'PATCH /incidents/:id/assign — Assign responder',
        'POST /incidents/:id/evidence — Upload evidence',
    ]),
    ('Patrols', [
        'GET/POST /patrol-teams — List and create patrol teams',
        'POST /patrol-teams/:id/members — Add team member',
        'GET/POST /patrol-shifts — List and create shifts',
        'GET /patrol-checkins — View check-in log',
    ]),
    ('Dashboard & Analytics', [
        'GET /dashboard/stats — Live stats for active patrols, alerts, incidents, SOS',
        'GET /dashboard/lga-summary — Per-LGA summary cards',
        'GET /analytics/incident-trends — Chart data for incident trends',
        'GET /analytics/response-times — Response time metrics',
        'GET /analytics/export — Export analytics (CSV/PDF)',
    ]),
    ('Communications', [
        'GET/POST /announcements — List and create broadcasts',
        'GET/POST /messages — Internal messaging',
        'PATCH /messages/:id/read — Mark message read',
    ]),
    ('Public (No Auth)', [
        'GET /health — API health check',
        'GET /alert-types — List alert type definitions',
        'GET /villages — List all villages (934+)',
        'GET /lgas — List all 9 LGAs',
        'GET /incidents — Active incidents for public map',
    ]),
    ('SMS/USSD', [
        'POST /sms/simulate — Simulate sending SMS (demo mode)',
        'GET  /sms/logs — View SMS delivery log',
    ]),
    ('Integration (API Key Auth)', [
        'POST /integration/sync/alerts — Sync alerts from layers',
        'POST /integration/sync/checkins — Sync patrol check-ins',
        'GET  /integration/config — Get integration configuration',
    ]),
]

for group_name, endpoints in endpoint_groups:
    add_para(doc, group_name, bold=True, color=(200, 16, 46), size=11)
    for ep in endpoints:
        add_bullet(doc, ep, level=1)
    doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 10 — SMS / USSD
# ═══════════════════════════════════════════════
add_heading(doc, '10. SMS / USSD Integration (*347#)', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc,
    'A critical feature distinguishing Ogbenjuwa from similar platforms is its full SMS/USSD '
    'fallback system. This allows citizens without smartphones or internet access to participate '
    'in the safety network using any basic phone.',
    size=11)

doc.add_paragraph()
add_para(doc, 'USSD Short Code: *347#', bold=True, size=13, color=(200, 16, 46))
doc.add_paragraph()

sms_commands = [
    ('PANIC [name] [village]', 'Send a panic/distress signal with location'),
    ('REPORT [type] [village]', 'Report an incident (type: attack/fire/medical/abduction)'),
    ('FIND ME [name] [age] [village]', 'Register in family reunification database'),
    ('FIND [name] [age] [LGA]', 'Search for a missing person by name'),
    ('*347*2*[LGA code]#', 'Find nearby resources (e.g., *347*2*AGT# for Agatu)'),
    ('*347*1#', 'View active alerts in your area'),
]

sms_table = doc.add_table(rows=1, cols=2)
sms_table.style = 'Table Grid'
hdr = sms_table.rows[0].cells
for i, h in enumerate(['Command', 'Action']):
    hdr[i].text = h
    set_cell_bg(hdr[i], 'C8102E')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

for i, (cmd, action) in enumerate(sms_commands):
    row = sms_table.add_row()
    row.cells[0].text = cmd
    row.cells[1].text = action
    bg = 'FFF5F5' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_para(doc, 'SMS Provider: Africa\'s Talking SDK  |  Delivery SLA: <30 seconds', size=10, italic=True, color=(100, 100, 100))
doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 11 — SECURITY
# ═══════════════════════════════════════════════
add_heading(doc, '11. Security Architecture', 1, color=(200, 16, 46))
add_divider(doc)

security_items = [
    ('JWT Authentication', 'Access tokens (15-min expiry, RS256) + Refresh tokens (7-day, httpOnly cookie). Auto-refresh on 401 response. Immediate invalidation on password change.'),
    ('Role-Based Access Control', 'Permission matrix enforced at middleware level. Route-level guards (requireRole), resource-level guards (requireScope), row-level LGA-scoped data isolation.'),
    ('API Key Authentication', 'SHA-256 hashed API keys for Layer 1/2 integration. Per-key rate limits, scope restrictions, usage tracking.'),
    ('Input Validation', 'Zod schema validation on all API inputs. TypeScript strict mode throughout — no any types.'),
    ('Rate Limiting', '5 login attempts per email per minute. Account lockout after 10 failed attempts. Per-endpoint and per-API-key rate limits.'),
    ('CORS', 'Restricted to known frontend origins. Environment-configured allowed origins.'),
    ('Audit Trail', 'Every administrative action logged: timestamp, user, action, resource ID, IP address.'),
    ('HTTPS', 'All production traffic encrypted via TLS. Vercel/Render enforce HTTPS.'),
    ('Multi-Factor Authentication', 'Optional TOTP-based 2FA for super_admin and lga_coordinator roles. Recovery codes (8 single-use codes).'),
]

for sec_title, sec_desc in security_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f'{sec_title}: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(sec_desc).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 12 — PERFORMANCE TARGETS
# ═══════════════════════════════════════════════
add_heading(doc, '12. Performance Targets', 1, color=(200, 16, 46))
add_divider(doc)

perf_table = doc.add_table(rows=1, cols=2)
perf_table.style = 'Table Grid'
hdr = perf_table.rows[0].cells
for i, h in enumerate(['Metric', 'Target']):
    hdr[i].text = h
    set_cell_bg(hdr[i], 'C8102E')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

perf_data = [
    ('API response time (p95)', '< 200ms'),
    ('Dashboard load time', '< 2 seconds'),
    ('WebSocket event latency', '< 100ms'),
    ('Concurrent API connections', '500+'),
    ('Concurrent WebSocket connections', '200+'),
    ('Database query time (p95)', '< 50ms'),
    ('Admin UI bundle size', '< 300KB gzipped'),
    ('Lighthouse score (admin UI)', '90+'),
    ('SMS delivery time', '< 30 seconds'),
    ('Page load on 3G connection', '< 5 seconds'),
    ('Mobile viewport support', 'From 360px width'),
]

for i, (metric, target) in enumerate(perf_data):
    row = perf_table.add_row()
    row.cells[0].text = metric
    row.cells[1].text = target
    bg = 'F5F5F5' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 13 — COVERAGE AREA
# ═══════════════════════════════════════════════
add_heading(doc, '13. Coverage Area — Idoma Region, Benue State', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc, 'Ogbenjuwa covers the 9 core Local Government Areas of Idoma Land:', size=11)
doc.add_paragraph()

lgas_table = doc.add_table(rows=1, cols=3)
lgas_table.style = 'Table Grid'
hdr = lgas_table.rows[0].cells
for i, h in enumerate(['LGA', 'Location in Idoma', 'Alert Zone']):
    hdr[i].text = h
    set_cell_bg(hdr[i], 'C8102E')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

lgas = [
    ('Otukpo', 'Central Idoma — Largest urban center', 'Zone A'),
    ('Ohimini', 'North-central Idoma', 'Zone A'),
    ('Okpokwu', 'Southern Idoma', 'Zone B'),
    ('Ado', 'Eastern Idoma', 'Zone B'),
    ('Ogbadibo', 'Western Idoma', 'Zone C'),
    ('Agatu', 'Riverine Idoma (Benue floodplain)', 'Zone C'),
    ('Apa', 'South-eastern Idoma', 'Zone D'),
    ('Obi', 'North-eastern Idoma', 'Zone D'),
    ('Oju', 'Northern Idoma', 'Zone E'),
]
for i, (lga, loc, zone) in enumerate(lgas):
    row = lgas_table.add_row()
    row.cells[0].text = lga
    row.cells[1].text = loc
    row.cells[2].text = zone
    bg = 'FFF5F5' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_para(doc, 'Seed Database Contains:', bold=True, size=11)
seed_stats = [
    '9 LGAs fully seeded', '24 wards mapped', '934+ villages with GPS coordinates',
    '5 alert types with SMS templates', '5 emergency contacts', '5 emergency resources',
    '5 family registry entries', '2 patrol teams', '3 announcements', '3 sample alerts',
    '1 default API key for integration',
]
for s in seed_stats:
    add_bullet(doc, s)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 14 — TEST ACCOUNTS
# ═══════════════════════════════════════════════
add_heading(doc, '14. Test Accounts & Demo Access', 1, color=(200, 16, 46))
add_divider(doc)

add_para(doc, 'All admin accounts share password: Password123!', bold=True, size=11)
doc.add_paragraph()

accounts_table = doc.add_table(rows=1, cols=4)
accounts_table.style = 'Table Grid'
hdr = accounts_table.rows[0].cells
for i, h in enumerate(['Email', 'Role', 'App URL', 'Access Level']):
    hdr[i].text = h
    set_cell_bg(hdr[i], '2D2D2D')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)

accounts_data = [
    ('daniel@ogbenjuwa.local', 'super_admin', 'localhost:4000', 'Full access — all features'),
    ('oche.agbo@ogbenjuwa.local', 'lga_coordinator', 'localhost:4000', 'LGA-level operations'),
    ('adah.ogiri@ogbenjuwa.local', 'vigilante_leader', 'localhost:4000', 'Patrol operations'),
    ('mama.ojoma@ogbenjuwa.local', 'community_admin', 'localhost:4000', 'Communications'),
    ('godwin.ibe@ogbenjuwa.local', 'resident', 'localhost:4000', 'Read-only'),
]
for i, (email, role, url, access) in enumerate(accounts_data):
    row = accounts_table.add_row()
    row.cells[0].text = email
    row.cells[1].text = role
    row.cells[2].text = url
    row.cells[3].text = access
    bg = 'F5F5F5' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
add_para(doc, 'Beacon Network & Citizen Portal:', bold=True, size=11)
add_bullet(doc, 'Login via phone number + OTP (demo: any 6-digit code is accepted)')
add_bullet(doc, 'Phone numbers: +2348034412290 (Daniel), +2348123456789 (Oche Agbo), +2348065432109 (Adah Ogiri), +2348157890123 (Mama Ojoma)')

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 15 — FUTURE ROADMAP
# ═══════════════════════════════════════════════
add_heading(doc, '15. Future Roadmap', 1, color=(200, 16, 46))
add_divider(doc)

roadmap = [
    ('Phase 2 — Mobile App', 'React Native app for iOS and Android, enabling push notifications and native GPS'),
    ('Phase 2 — Real-time PWA', 'Service Worker background sync for offline reports; Web Push API for instant alerts'),
    ('Phase 2 — WebSocket Full Rollout', 'Real-time alert propagation to all Beacon Network users without page refresh'),
    ('Phase 3 — Multi-language', 'Additional language support: Hausa, Igbo, Pidgin English, Tiv for border communities'),
    ('Phase 3 — Herdsmen Route Tracking', 'Predictive mapping of cattle migration routes to proactively alert communities'),
    ('Phase 3 — AI Incident Classification', 'Machine learning to auto-classify incoming SMS reports and prioritize response'),
    ('Phase 4 — State Government Integration', 'Direct API integration with Benue State Emergency Management Agency (SEMA)'),
    ('Phase 4 — Revenue Model', 'Government subscriptions (LGA-level), NGO grants, diaspora community crowdfunding'),
]

for phase, detail in roadmap:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f'{phase}: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(detail).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# SECTION 16 — TECH STACK SUMMARY
# ═══════════════════════════════════════════════
add_heading(doc, '16. Complete Technology Stack', 1, color=(200, 16, 46))
add_divider(doc)

tech_table = doc.add_table(rows=1, cols=3)
tech_table.style = 'Table Grid'
hdr = tech_table.rows[0].cells
for i, h in enumerate(['Category', 'Technology', 'Purpose']):
    hdr[i].text = h
    set_cell_bg(hdr[i], 'C8102E')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

tech_data = [
    ('Frontend Framework', 'React 18 + TypeScript', 'UI components and state management'),
    ('Build Tool', 'Vite 5', 'Fast development server and production builds'),
    ('UI Components', 'Shadcn UI + Radix UI', 'Accessible, composable component primitives'),
    ('Styling', 'Tailwind CSS 3', 'Utility-first responsive styling'),
    ('Animations', 'Framer Motion', 'Smooth transitions and micro-animations'),
    ('Maps', 'Leaflet + react-leaflet', 'Interactive alert and patrol maps'),
    ('Icons', 'Lucide React', 'Consistent icon library throughout'),
    ('HTTP Client', 'Axios', 'API communication with JWT interceptors'),
    ('Backend', 'Node.js + Express', 'REST API server and WebSocket host'),
    ('Language', 'TypeScript (strict)', 'Type safety with Zod runtime validation'),
    ('ORM', 'Drizzle ORM', 'Type-safe PostgreSQL queries and migrations'),
    ('Database', 'PostgreSQL (Neon)', 'Serverless PostgreSQL with connection pooling'),
    ('Authentication', 'JWT (jsonwebtoken)', 'Stateless auth with refresh token rotation'),
    ('Password Hashing', 'bcrypt (12 rounds)', 'Secure credential storage'),
    ('SMS/USSD', "Africa's Talking SDK", 'SMS delivery and USSD menu system'),
    ('Real-time', 'WebSocket (ws)', 'Live dashboard updates and alert propagation'),
    ('Package Manager', 'Bun', 'Fast JS runtime and package manager'),
    ('Container', 'Docker', 'Production containerization on Render'),
    ('Frontend Deploy', 'Vercel', 'CDN-deployed static sites with CI/CD'),
    ('Backend Deploy', 'Render', 'Docker web service for the API'),
    ('Database Host', 'Neon / Supabase', 'Managed PostgreSQL with branching'),
    ('Version Control', 'Git + GitHub', 'Source control and CI/CD integration'),
]

for i, (cat, tech, purpose) in enumerate(tech_data):
    row = tech_table.add_row()
    row.cells[0].text = cat
    row.cells[1].text = tech
    row.cells[2].text = purpose
    bg = 'FFF5F5' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# ═══════════════════════════════════════════════
# CLOSING
# ═══════════════════════════════════════════════
doc.add_page_break()

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run('Ogbenjuwa Community Safety Network')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(200, 16, 46)

closing2 = doc.add_paragraph()
closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing2.add_run('Built with purpose. Protecting Idoma lives.').font.size = Pt(13)

doc.add_paragraph()

contact_info = doc.add_paragraph()
contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_info.add_run(
    'GitHub: github.com/AmPhilDanny/Ogbenjuwa-beacon-network\n'
    'Website: ogbenjuwa.com.ng\n'
    'Portal: portal.ogbenjuwa.com.ng\n'
    'Admin: central-command.ogbenjuwa.com.ng\n'
    'Emergency SMS: *347#'
)
for run in contact_info.runs:
    run.font.size = Pt(11)

doc.add_paragraph()
gen_date = doc.add_paragraph()
gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
gen_date.add_run(f'Document generated: {datetime.datetime.now().strftime("%B %d, %Y")}').font.size = Pt(10)

# ─────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────
output_path = r'C:\Users\user pc\Desktop\AMUA PROJECT\Ogbenjuwa_Project_Documentation.docx'
doc.save(output_path)
print(f'✅ Document saved: {output_path}')
print(f'   Pages: ~30+')
print(f'   Sections: 16 major sections')
